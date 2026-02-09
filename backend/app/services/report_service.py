"""
周报生成服务层
提供周报和工作日志相关的业务逻辑实现
"""

from typing import List, Optional, Dict, Any
import asyncio
from datetime import datetime, timedelta
import os
from sqlalchemy import select, and_, or_, func
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models.report import WeeklyReport, WorkLog, ReportStatus
from app.schemas.report import (
    WorkLogCreate, WorkLogUpdate, WorkLogResponse,
    WeeklyReportCreate, WeeklyReportUpdate, WeeklyReportReview,
    WeeklyReportResponse, WeeklyReportDetailResponse
)
from app.utils.logger import get_logger
from app.utils.exceptions import ValidationError
from app.core.config import settings
from app.services.llm_service import llm_service

logger = get_logger(__name__)


class WeeklyReportService:
    """周报生成服务 - 处理周报和工作日志相关业务逻辑"""
    
    def __init__(self, db: AsyncSession):
        self.db = db
    
    # ============================================================
    # 工作日志相关方法
    # ============================================================
    
    async def create_log(self, log_data: WorkLogCreate, user_id: Optional[int] = None) -> WorkLogResponse:
        """
        创建工作日志
        
        Args:
            log_data: 工作日志数据
            user_id: 用户ID（可选）
            
        Returns:
            创建的工作日志信息
        """
        try:
            logger.info(f"创建工作日志: {log_data.work_type}")
            
            # 确定日志日期
            log_date = log_data.log_date or datetime.utcnow()
            
            # 创建日志记录
            work_log = WorkLog(
                user_id=user_id,
                work_type=log_data.work_type,
                task_description=log_data.task_description,
                hours_spent=log_data.hours_spent,
                log_date=log_date
            )
            
            self.db.add(work_log)
            await self.db.commit()
            await self.db.refresh(work_log)
            
            logger.info(f"工作日志创建成功: {work_log.id}")
            return WorkLogResponse.model_validate(work_log)
        
        except Exception as e:
            await self.db.rollback()
            logger.error(f"创建工作日志失败: {e}")
            raise
    
    async def list_logs(
        self,
        user_id: Optional[int] = None,
        date_from: Optional[datetime] = None,
        date_to: Optional[datetime] = None,
        skip: int = 0,
        limit: int = 100
    ) -> Dict:
        """
        获取工作日志列表
        
        Args:
            user_id: 用户ID（可选）
            date_from: 起始日期
            date_to: 结束日期
            skip: 分页偏移
            limit: 分页大小
            
        Returns:
            日志列表及分页信息
        """
        try:
            # 构建查询条件
            conditions = []
            if user_id is not None:
                conditions.append(WorkLog.user_id == user_id)
            if date_from:
                conditions.append(WorkLog.log_date >= date_from)
            if date_to:
                conditions.append(WorkLog.log_date <= date_to)
            
            # 查询总数
            count_stmt = select(func.count()).select_from(WorkLog)
            if conditions:
                count_stmt = count_stmt.where(and_(*conditions))
            count_result = await self.db.execute(count_stmt)
            total = count_result.scalar_one()
            
            # 执行查询
            query = select(WorkLog)
            if conditions:
                query = query.where(and_(*conditions))
            query = query.order_by(WorkLog.log_date.desc()).offset(skip).limit(limit)
            
            result = await self.db.execute(query)
            logs = result.scalars().all()
            
            logger.info(f"查询工作日志列表: 总数={total}")
            
            return {
                "total": total,
                "skip": skip,
                "limit": limit,
                "items": [WorkLogResponse.model_validate(log) for log in logs]
            }
        
        except Exception as e:
            logger.error(f"查询工作日志列表失败: {e}")
            raise
    
    async def get_log(self, log_id: int, user_id: Optional[int] = None) -> WorkLogResponse:
        """获取工作日志详情"""
        try:
            conditions = [WorkLog.id == log_id]
            if user_id is not None:
                conditions.append(WorkLog.user_id == user_id)
            query = select(WorkLog).where(and_(*conditions))
            result = await self.db.execute(query)
            log = result.scalar_one_or_none()
            
            if not log:
                raise ValidationError(f"工作日志不存在: {log_id}")
            
            return WorkLogResponse.model_validate(log)
        
        except Exception as e:
            logger.error(f"获取工作日志失败: {e}")
            raise
    
    async def update_log(
        self, log_id: int, log_data: WorkLogUpdate, user_id: Optional[int] = None
    ) -> WorkLogResponse:
        """更新工作日志"""
        try:
            conditions = [WorkLog.id == log_id]
            if user_id is not None:
                conditions.append(WorkLog.user_id == user_id)
            query = select(WorkLog).where(and_(*conditions))
            result = await self.db.execute(query)
            log = result.scalar_one_or_none()
            
            if not log:
                raise ValidationError(f"工作日志不存在: {log_id}")
            
            # 更新字段
            if log_data.work_type is not None:
                log.work_type = log_data.work_type
            if log_data.task_description is not None:
                log.task_description = log_data.task_description
            if log_data.hours_spent is not None:
                log.hours_spent = log_data.hours_spent
            
            log.updated_at = datetime.utcnow()
            
            await self.db.commit()
            await self.db.refresh(log)
            
            logger.info(f"工作日志更新成功: {log_id}")
            return WorkLogResponse.model_validate(log)
        
        except Exception as e:
            await self.db.rollback()
            logger.error(f"更新工作日志失败: {e}")
            raise
    
    async def delete_log(self, log_id: int, user_id: Optional[int] = None) -> None:
        """删除工作日志"""
        try:
            conditions = [WorkLog.id == log_id]
            if user_id is not None:
                conditions.append(WorkLog.user_id == user_id)
            query = select(WorkLog).where(and_(*conditions))
            result = await self.db.execute(query)
            log = result.scalar_one_or_none()
            
            if not log:
                raise ValidationError(f"工作日志不存在: {log_id}")
            
            await self.db.delete(log)
            await self.db.commit()
            
            logger.info(f"工作日志删除成功: {log_id}")
        
        except Exception as e:
            await self.db.rollback()
            logger.error(f"删除工作日志失败: {e}")
            raise
    
    # ============================================================
    # 周报相关方法
    # ============================================================
    
    def _generate_week_identifier(self, date: datetime) -> str:
        """生成周标识符 (如 2025-W04)"""
        year, week, _ = date.isocalendar()
        return f"{year}-W{week:02d}"
    
    async def create_report(
        self,
        report_data: WeeklyReportCreate,
        user_id: Optional[int] = None
    ) -> WeeklyReportDetailResponse:
        """
        生成周报
        
        Args:
            report_data: 周报数据
            user_id: 用户ID（可选）
            
        Returns:
            创建的周报信息
        """
        try:
            # 验证日期范围
            if report_data.week_end_date <= report_data.week_start_date:
                raise ValidationError("周结束日期必须晚于周开始日期")
            
            # 生成周标识符
            week_identifier = self._generate_week_identifier(report_data.week_start_date)
            
            logger.info(f"生成周报: {week_identifier}")
            
            # 获取该周的工作日志并计算总工时
            logs_query = select(WorkLog).where(
                and_(
                    WorkLog.log_date >= report_data.week_start_date,
                    WorkLog.log_date <= report_data.week_end_date,
                    WorkLog.user_id == user_id
                )
            )
            logs_result = await self.db.execute(logs_query)
            logs = logs_result.scalars().all()
            
            total_hours = sum(log.hours_spent for log in logs)
            
            # 生成周报摘要与详细内容
            summary = self._generate_summary(logs)
            content = self._generate_detail(logs, week_identifier)

            # AI 扩写润色
            if report_data.ai_polish and llm_service.check_availability():
                loop = asyncio.get_running_loop()
                polished = await loop.run_in_executor(
                    None,
                    llm_service.polish_weekly_report,
                    summary,
                    content
                )
                summary = polished.get("summary", summary)
                content = polished.get("content", content)
            
            # 创建周报记录
            report = WeeklyReport(
                user_id=user_id,
                title=report_data.title or f"周报 {week_identifier}",
                week_start_date=report_data.week_start_date,
                week_end_date=report_data.week_end_date,
                week=week_identifier,
                summary=summary,
                content=content,
                total_hours=total_hours,
                status=ReportStatus.DRAFT
            )
            
            self.db.add(report)
            await self.db.commit()
            await self.db.refresh(report)
            
            logger.info(f"周报生成成功: {report.id} ({week_identifier})")
            return WeeklyReportDetailResponse.model_validate(report)
        
        except Exception as e:
            await self.db.rollback()
            logger.error(f"生成周报失败: {e}")
            raise
    
    def _generate_summary(self, logs: List[WorkLog]) -> str:
        """根据工作日志生成周报摘要"""
        if not logs:
            return "本周完成：\n- 本周暂无工作记录"

        grouped = self._group_logs(logs)
        top_tasks = self._pick_top_tasks(grouped, limit=5)
        risks = self._extract_risks(logs)
        plans = self._extract_plans(logs)

        lines = ["本周完成："]
        if top_tasks:
            lines.extend([f"- {item}" for item in top_tasks])
        else:
            lines.append("- 完成常规工作推进")

        lines.append("")
        lines.append("问题与风险：")
        if risks:
            lines.extend([f"- {item}" for item in risks])
        else:
            lines.append("- 暂无突出问题")

        lines.append("")
        lines.append("下周计划：")
        if plans:
            lines.extend([f"- {item}" for item in plans])
        else:
            lines.append("- 继续推进本周任务并跟进结果")

        return "\n".join(lines)

    def _generate_detail(self, logs: List[WorkLog], week_identifier: str) -> str:
        """生成更清晰的周报详细内容"""
        if not logs:
            return "本周暂无工作记录。"

        grouped = self._group_logs(logs)

        lines = [f"# 周报详情 ({week_identifier})", "", "## 工作概览"]
        for work_type, data in grouped.items():
            lines.append(f"- {work_type}: {data['hours']:.1f}小时")

        lines.append("")
        lines.append("## 关键进展")
        for work_type, data in grouped.items():
            lines.append(f"### {work_type}")
            for task in data["tasks"]:
                lines.append(f"- {task}")
            lines.append("")

        risks = self._extract_risks(logs)
        lines.append("## 问题与风险")
        if risks:
            lines.extend([f"- {item}" for item in risks])
        else:
            lines.append("- 暂无突出问题")

        plans = self._extract_plans(logs)
        lines.append("")
        lines.append("## 下周计划")
        if plans:
            lines.extend([f"- {item}" for item in plans])
        else:
            lines.append("- 继续推进本周任务并跟进结果")

        return "\n".join(lines)

    def _group_logs(self, logs: List[WorkLog]) -> Dict[str, Dict[str, Any]]:
        grouped: Dict[str, Dict[str, Any]] = {}
        for log in logs:
            work_type = log.work_type or "其他"
            if work_type not in grouped:
                grouped[work_type] = {"hours": 0.0, "tasks": []}
            grouped[work_type]["hours"] += float(log.hours_spent or 0)
            grouped[work_type]["tasks"].append(log.task_description)
        return grouped

    def _pick_top_tasks(self, grouped: Dict[str, Dict[str, Any]], limit: int = 5) -> List[str]:
        tasks: List[str] = []
        for work_type, data in grouped.items():
            for task in data["tasks"]:
                tasks.append(f"{work_type}：{task}")
        return tasks[:limit]

    def _extract_risks(self, logs: List[WorkLog]) -> List[str]:
        keywords = ["风险", "问题", "阻塞", "延期", "卡点", "Bug", "故障", "缺陷", "异常"]
        results = []
        for log in logs:
            desc = log.task_description or ""
            if any(k.lower() in desc.lower() for k in keywords):
                results.append(desc)
        return results[:5]

    def _extract_plans(self, logs: List[WorkLog]) -> List[str]:
        keywords = ["下周", "计划", "准备", "安排", "目标", "排期", "优化", "跟进"]
        results = []
        for log in logs:
            desc = log.task_description or ""
            if any(k in desc for k in keywords):
                results.append(desc)
        return results[:5]
    
    async def list_reports(
        self,
        user_id: Optional[int] = None,
        status: Optional[str] = None,
        skip: int = 0,
        limit: int = 10
    ) -> Dict:
        """
        获取周报列表
        
        Args:
            user_id: 用户ID（可选）
            status: 周报状态过滤
            skip: 分页偏移
            limit: 分页大小
            
        Returns:
            周报列表及分页信息
        """
        try:
            conditions = []
            if user_id is not None:
                conditions.append(WeeklyReport.user_id == user_id)
            if status:
                conditions.append(WeeklyReport.status == status)
            
            # 查询总数
            count_query = select(WeeklyReport)
            if conditions:
                count_query = count_query.where(and_(*conditions))
            count_result = await self.db.execute(count_query)
            total = len(count_result.scalars().all())
            
            # 执行查询
            query = select(WeeklyReport)
            if conditions:
                query = query.where(and_(*conditions))
            query = query.order_by(WeeklyReport.created_at.desc()).offset(skip).limit(limit)
            
            result = await self.db.execute(query)
            reports = result.scalars().all()
            
            logger.info(f"查询周报列表: 总数={total}")
            
            return {
                "total": total,
                "skip": skip,
                "limit": limit,
                "items": [WeeklyReportResponse.model_validate(report) for report in reports]
            }
        
        except Exception as e:
            logger.error(f"查询周报列表失败: {e}")
            raise
    
    async def get_report(
        self, report_id: int, user_id: Optional[int] = None
    ) -> WeeklyReportDetailResponse:
        """获取周报详情"""
        try:
            conditions = [WeeklyReport.id == report_id]
            if user_id is not None:
                conditions.append(WeeklyReport.user_id == user_id)
            query = select(WeeklyReport).where(and_(*conditions))
            result = await self.db.execute(query)
            report = result.scalar_one_or_none()
            
            if not report:
                raise ValidationError(f"周报不存在: {report_id}")
            
            return WeeklyReportDetailResponse.model_validate(report)
        
        except Exception as e:
            logger.error(f"获取周报详情失败: {e}")
            raise
    
    async def update_report(
        self,
        report_id: int,
        report_data: WeeklyReportUpdate,
        user_id: Optional[int] = None
    ) -> WeeklyReportDetailResponse:
        """更新周报"""
        try:
            conditions = [WeeklyReport.id == report_id]
            if user_id is not None:
                conditions.append(WeeklyReport.user_id == user_id)
            query = select(WeeklyReport).where(and_(*conditions))
            result = await self.db.execute(query)
            report = result.scalar_one_or_none()
            
            if not report:
                raise ValidationError(f"周报不存在: {report_id}")
            
            # 只允许编辑草稿状态的周报
            if report.status != ReportStatus.DRAFT:
                raise ValidationError(f"只能编辑草稿状态的周报")
            
            # 更新字段
            if report_data.title is not None:
                report.title = report_data.title
            if report_data.summary is not None:
                report.summary = report_data.summary
            if report_data.content is not None:
                report.content = report_data.content
            
            report.updated_at = datetime.utcnow()
            
            await self.db.commit()
            await self.db.refresh(report)
            
            logger.info(f"周报更新成功: {report_id}")
            return WeeklyReportDetailResponse.model_validate(report)
        
        except Exception as e:
            await self.db.rollback()
            logger.error(f"更新周报失败: {e}")
            raise
    
    async def delete_report(self, report_id: int, user_id: Optional[int] = None) -> None:
        """删除周报（只允许删除草稿）"""
        try:
            conditions = [WeeklyReport.id == report_id]
            if user_id is not None:
                conditions.append(WeeklyReport.user_id == user_id)
            query = select(WeeklyReport).where(and_(*conditions))
            result = await self.db.execute(query)
            report = result.scalar_one_or_none()
            
            if not report:
                raise ValidationError(f"周报不存在: {report_id}")
            
            if report.status != ReportStatus.DRAFT:
                raise ValidationError(f"只能删除草稿状态的周报")
            
            await self.db.delete(report)
            await self.db.commit()
            
            logger.info(f"周报删除成功: {report_id}")
        
        except Exception as e:
            await self.db.rollback()
            logger.error(f"删除周报失败: {e}")
            raise
    
    async def submit_report(
        self, report_id: int, user_id: Optional[int] = None
    ) -> WeeklyReportDetailResponse:
        """提交周报审核"""
        try:
            conditions = [WeeklyReport.id == report_id]
            if user_id is not None:
                conditions.append(WeeklyReport.user_id == user_id)
            query = select(WeeklyReport).where(and_(*conditions))
            result = await self.db.execute(query)
            report = result.scalar_one_or_none()
            
            if not report:
                raise ValidationError(f"周报不存在: {report_id}")
            
            if report.status != ReportStatus.DRAFT:
                raise ValidationError(f"只能提交草稿状态的周报")
            
            report.status = ReportStatus.SUBMITTED
            report.updated_at = datetime.utcnow()
            
            await self.db.commit()
            await self.db.refresh(report)
            
            logger.info(f"周报提交成功: {report_id}")
            return WeeklyReportDetailResponse.model_validate(report)
        
        except Exception as e:
            await self.db.rollback()
            logger.error(f"提交周报失败: {e}")
            raise
    
    async def review_report(
        self,
        report_id: int,
        review_data: WeeklyReportReview,
        reviewer_id: Optional[int] = None,
        user_id: Optional[int] = None
    ) -> WeeklyReportDetailResponse:
        """审核周报"""
        try:
            conditions = [WeeklyReport.id == report_id]
            if user_id is not None:
                conditions.append(WeeklyReport.user_id == user_id)
            query = select(WeeklyReport).where(and_(*conditions))
            result = await self.db.execute(query)
            report = result.scalar_one_or_none()
            
            if not report:
                raise ValidationError(f"周报不存在: {report_id}")
            
            if report.status != ReportStatus.SUBMITTED:
                raise ValidationError(f"只能审核已提交的周报")
            
            # 验证审核状态
            if review_data.status not in ["approved", "rejected"]:
                raise ValidationError(f"无效的审核状态: {review_data.status}")
            
            report.status = ReportStatus.APPROVED if review_data.status == "approved" else ReportStatus.REJECTED
            report.review_feedback = review_data.review_feedback
            report.reviewer_id = reviewer_id
            report.reviewed_at = datetime.utcnow()
            report.updated_at = datetime.utcnow()
            
            await self.db.commit()
            await self.db.refresh(report)
            
            logger.info(f"周报审核完成: {report_id}, 状态={review_data.status}")
            return WeeklyReportDetailResponse.model_validate(report)
        
        except Exception as e:
            await self.db.rollback()
            logger.error(f"审核周报失败: {e}")
            raise
    
    async def export_report(
        self, report_id: int, format: str = "markdown", user_id: Optional[int] = None
    ) -> Dict:
        """
        导出周报
        
        Args:
            report_id: 周报ID
            format: 导出格式 (markdown/html)
            
        Returns:
            导出的周报内容
        """
        try:
            conditions = [WeeklyReport.id == report_id]
            if user_id is not None:
                conditions.append(WeeklyReport.user_id == user_id)
            query = select(WeeklyReport).where(and_(*conditions))
            result = await self.db.execute(query)
            report = result.scalar_one_or_none()
            
            if not report:
                raise ValidationError(f"周报不存在: {report_id}")
            if format == "markdown":
                content = self._export_as_markdown(report)
                logger.info(f"周报导出成功: {report_id}, 格式={format}")
                return {"report_id": report_id, "format": format, "content": content}
            if format == "html":
                content = self._export_as_html(report)
                logger.info(f"周报导出成功: {report_id}, 格式={format}")
                return {"report_id": report_id, "format": format, "content": content}
            if format in {"pdf", "docx"}:
                path = self._export_as_document(report, format)
                logger.info(f"周报导出成功: {report_id}, 格式={format}")
                return {"report_id": report_id, "format": format, "path": path}

            raise ValidationError(f"不支持的导出格式: {format}")
        
        except Exception as e:
            logger.error(f"导出周报失败: {e}")
            raise
    
    def _export_as_markdown(self, report: WeeklyReport) -> str:
        """导出为Markdown格式"""
        lines = [
            f"# {report.title}",
            f"\n**周期**: {report.week_start_date.strftime('%Y-%m-%d')} 至 {report.week_end_date.strftime('%Y-%m-%d')}",
            f"\n**总工时**: {report.total_hours:.1f}小时",
            f"\n**状态**: {report.status.value}",
        ]
        
        if report.summary:
            lines.extend(["\n## 摘要\n", report.summary])
        
        if report.content:
            lines.extend(["\n## 详细内容\n", report.content])
        
        if report.review_feedback:
            lines.extend(["\n## 审核反馈\n", report.review_feedback])
        
        return "\n".join(lines)
    
    def _export_as_html(self, report: WeeklyReport) -> str:
        """导出为HTML格式"""
        html_parts = [
            f"<h1>{report.title}</h1>",
            f"<p><strong>周期</strong>: {report.week_start_date.strftime('%Y-%m-%d')} 至 {report.week_end_date.strftime('%Y-%m-%d')}</p>",
            f"<p><strong>总工时</strong>: {report.total_hours:.1f}小时</p>",
            f"<p><strong>状态</strong>: {report.status.value}</p>",
        ]

        if report.summary:
            html_parts.extend([
                "<h2>摘要</h2>",
                f"<p>{report.summary.replace(chr(10), '<br>')}</p>"
            ])

        if report.content:
            html_parts.extend([
                "<h2>详细内容</h2>",
                f"<p>{report.content.replace(chr(10), '<br>')}</p>"
            ])

        if report.review_feedback:
            html_parts.extend([
                "<h2>审核反馈</h2>",
                f"<p>{report.review_feedback.replace(chr(10), '<br>')}</p>"
            ])

        return "".join(html_parts)

    def _export_as_document(self, report: WeeklyReport, format: str) -> str:
        """导出为PDF或DOCX文件"""
        upload_dir = settings.UPLOAD_DIR
        os.makedirs(upload_dir, exist_ok=True)

        filename = f"weekly_report_{report.id}_{report.week}.{format}"
        file_path = os.path.join(upload_dir, filename)

        if format == "pdf":
            success = self._generate_weekly_pdf(report, file_path)
        else:
            success = self._generate_weekly_docx(report, file_path)

        if not success:
            raise ValidationError(f"{format.upper()} 生成失败，请检查依赖是否安装")

        return f"/uploads/{filename}"

    def _generate_weekly_pdf(self, report: WeeklyReport, output_path: str) -> bool:
        """生成周报 PDF"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            import platform

            font_registered = False
            system = platform.system()
            font_name = "SimHei"
            font_paths: List[str] = []
            if system == "Windows":
                font_paths = [
                    "C:\\Windows\\Fonts\\simhei.ttf",
                    "C:\\Windows\\Fonts\\SimHei.ttf",
                    "C:\\Windows\\Fonts\\msyh.ttf",
                    "C:\\Windows\\Fonts\\msyhbd.ttf",
                ]
            elif system == "Darwin":
                font_paths = [
                    "/Library/Fonts/SimHei.ttf",
                    "/System/Library/Fonts/STHeiti Medium.ttc",
                    "/System/Library/Fonts/PingFang.ttc",
                ]
            else:
                font_paths = [
                    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                    "/usr/share/fonts/wqy-microhei/wqy-microhei.ttc",
                ]

            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        font_registered = True
                        break
                    except Exception:
                        continue

            if not font_registered:
                try:
                    font_name = "STSong-Light"
                    pdfmetrics.registerFont(UnicodeCIDFont(font_name))
                    font_registered = True
                except Exception:
                    font_registered = False

            styles = getSampleStyleSheet()
            base_style = styles["BodyText"]
            if font_registered:
                base_style.fontName = font_name

            title_style = ParagraphStyle(
                "WeeklyTitle",
                parent=styles["Title"],
                fontName=font_name if font_registered else styles["Title"].fontName,
            )
            heading_style = ParagraphStyle(
                "WeeklyHeading",
                parent=styles["Heading2"],
                fontName=font_name if font_registered else styles["Heading2"].fontName,
            )

            doc = SimpleDocTemplate(output_path, pagesize=A4)
            elements = []
            title = report.title or f"周报 {report.week}"
            date_range = f"{report.week_start_date.strftime('%Y-%m-%d')} 至 {report.week_end_date.strftime('%Y-%m-%d')}"

            elements.append(Paragraph(title, title_style))
            elements.append(Spacer(1, 12))
            elements.append(Paragraph(f"周期：{date_range}", base_style))
            elements.append(Paragraph(f"总工时：{report.total_hours:.1f}小时", base_style))
            elements.append(Paragraph(f"状态：{report.status.value}", base_style))
            elements.append(Spacer(1, 12))

            if report.summary:
                elements.append(Paragraph("摘要", heading_style))
                elements.append(Paragraph(report.summary.replace("\n", "<br/>") or "-", base_style))
                elements.append(Spacer(1, 12))

            if report.content:
                elements.append(Paragraph("详细内容", heading_style))
                elements.append(Paragraph(report.content.replace("\n", "<br/>") or "-", base_style))
                elements.append(Spacer(1, 12))

            if report.review_feedback:
                elements.append(Paragraph("审核反馈", heading_style))
                elements.append(Paragraph(report.review_feedback.replace("\n", "<br/>") or "-", base_style))

            doc.build(elements)
            logger.info(f"周报 PDF 生成成功: {output_path}")
            return True
        except ImportError as e:
            logger.error(f"reportlab 未安装: {e}")
            return False
        except Exception as e:
            logger.error(f"周报 PDF 生成失败: {type(e).__name__}: {e}", exc_info=True)
            return False

    def _generate_weekly_docx(self, report: WeeklyReport, output_path: str) -> bool:
        """生成周报 Word"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "宋体"
            style.font.size = Pt(12)

            title = report.title or f"周报 {report.week}"
            date_range = f"{report.week_start_date.strftime('%Y-%m-%d')} 至 {report.week_end_date.strftime('%Y-%m-%d')}"

            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"周期：{date_range}")
            doc.add_paragraph(f"总工时：{report.total_hours:.1f}小时")
            doc.add_paragraph(f"状态：{report.status.value}")

            if report.summary:
                doc.add_heading("摘要", level=2)
                for line in report.summary.splitlines():
                    doc.add_paragraph(line)

            if report.content:
                doc.add_heading("详细内容", level=2)
                for line in report.content.splitlines():
                    doc.add_paragraph(line)

            if report.review_feedback:
                doc.add_heading("审核反馈", level=2)
                for line in report.review_feedback.splitlines():
                    doc.add_paragraph(line)

            doc.save(output_path)
            logger.info(f"周报 Word 生成成功: {output_path}")
            return True
        except ImportError:
            logger.error("python-docx 未安装，请运行: pip install python-docx")
            return False
        except Exception as e:
            logger.error(f"周报 Word 生成失败: {type(e).__name__}: {e}", exc_info=True)
            return False
