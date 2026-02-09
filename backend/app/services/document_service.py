"""
文献摘要服务层
提供文档处理和摘要生成逻辑
"""

from typing import List, Optional, Dict, Any
from datetime import datetime
import json
import os
import re
import uuid

import aiohttp
import pdfplumber
from docx import Document as DocxDocument
from fastapi import UploadFile
from sqlalchemy import select, update, delete, and_
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.document import Document, DocumentSummary
from app.models.knowledge import DocumentVersion, Review, Tag
from app.services.search_service import search_service
from sqlalchemy.orm import selectinload
from app.utils.logger import get_logger
from app.services.weknora_service import weknora_service
from app.services.llm_service import llm_service
from app.services.chunking_service import chunking_service

logger = get_logger(__name__)


class DocumentService:
    """文献文档服务"""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.weknora = weknora_service

    async def _get_or_create_default_kb(self, user_id: int) -> str:
        """获取或为用户创建默认知识库"""
        # 这里简单处理，先尝试找现有的知识库，没有则创建一个名为 "Default" 的
        # 在实际生产中，可能需要一个单独的表来记录用户的知识库映射
        try:
            kbs = await self.weknora._request("GET", "/knowledge-bases")
            for kb in kbs.get("data", []):
                if kb.get("name") == f"User_{user_id}_Default":
                    return kb.get("id")
            
            # 没找到，创建一个
            # 必须指定 Embedding 模型，否则无法进行 RAG 问答
            # 同时绑定 LLM 模型用于后续的摘要和问答优化
            embedding_model_id = await self.weknora.get_embedding_model_id("text-embedding-v3")
            chat_model_id = await self.weknora.get_model_id_by_type("Chat")
            
            new_kb = await self.weknora.create_knowledge_base(
                name=f"User_{user_id}_Default", 
                description=f"Default knowledge base for user {user_id}",
                embedding_model_id=embedding_model_id,
                summary_model_id=chat_model_id
            )
            return new_kb.get("data", {}).get("id")
        except Exception as e:
            logger.error(f"获取或创建 WeKnora 知识库失败: {e}")
            return None

    async def create_document(self, title: str, file: UploadFile, user_id: int, kb_id: Optional[int] = None, dir_id: Optional[int] = None) -> dict:
        """
        创建文档 (第一阶段：保存文件并创建记录)
        """
        logger.info(f"开始创建文档任务: {title}")

        filename = file.filename or "document"
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in {"pdf", "txt", "docx", "md"}:
            raise ValueError("仅支持 PDF/TXT/DOCX/MD 文件")

        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        safe_name = f"doc_{uuid.uuid4().hex}.{ext}"
        save_path = os.path.join(settings.UPLOAD_DIR, safe_name)

        content = await file.read()
        if not content:
            raise ValueError("文件内容为空")

        with open(save_path, "wb") as f:
            f.write(content)

        # 初始化文档记录
        meta_info = {
            "filename": filename,
            "size": len(content),
            "ext": ext,
        }

        # 预先获取或创建默认 WeKnora 知识库 ID (如果没提供本地 kb_id)
        # 移除 WeKnora 自动创建逻辑
        weknora_kb_id = None
        # if not kb_id:
        #      weknora_kb_id = await self._get_or_create_default_kb(user_id)

        doc = Document(
            user_id=user_id,
            title=title or os.path.splitext(filename)[0],
            content="",  # 暂时为空，异步处理时填充
            document_type="source",
            source_type="file",
            file_path=save_path,
            meta_info=json.dumps(meta_info, ensure_ascii=False),
            knowledge_base_id=kb_id,
            directory_id=dir_id,
            weknora_kb_id=weknora_kb_id,
            status="pending",
            review_status="draft",
            processing_progress=0
        )
        self.db.add(doc)
        await self.db.commit()
        await self.db.refresh(doc)

        # 创建初始版本
        version = DocumentVersion(
            document_id=doc.id,
            version_number=1,
            title=doc.title,
            content="",
            created_by=user_id,
            change_log="初始上传"
        )
        self.db.add(version)
        await self.db.commit()

        return self._format_document(doc)

    async def process_document_background(self, doc_id: int):
        """
        后台处理文档 (解析 + 同步到 WeKnora)
        """
        logger.info(f"开始后台处理文档: {doc_id}")
        
        # 获取最新的 session 进行操作
        async with self.db as session:
            try:
                # 重新查询文档
                query = select(Document).where(Document.id == doc_id)
                result = await session.execute(query)
                doc = result.scalars().first()
                
                if not doc:
                    logger.error(f"文档 {doc_id} 不存在，停止处理")
                    return

                # 更新状态：处理中
                doc.status = "processing"
                doc.processing_progress = 10
                await session.commit()

                # 1. 解析文件内容
                parsed_text = ""
                try:
                    ext = json.loads(doc.meta_info).get("ext", "")
                    parsed_text = self._parse_file(doc.file_path, ext)
                    if not parsed_text.strip():
                        # 对于空内容或者无法解析的情况，不应该直接报错停止，
                        # 而是记录警告，允许用户后续手动编辑内容。
                        logger.warning(f"文件解析为空: {doc.filename}")
                        # raise ValueError("未解析到有效文本") 
                    
                    doc.content = parsed_text
                    doc.processing_progress = 40
                    await session.commit()
                except Exception as e:
                    logger.error(f"解析文档失败: {e}")
                    # 解析失败不应完全终止流程，可以保留空内容供后续编辑
                    doc.status = "warning" 
                    doc.error_message = f"解析警告: {str(e)}"
                    doc.content = ""
                    await session.commit()
                    # 继续执行后续流程，至少让文档记录存在并可搜索（标题等）

                # 2. 同步到 WeKnora (已禁用)
                # if doc.weknora_kb_id:
                #     try:
                #         weknora_result = await self.weknora.upload_document_file(doc.weknora_kb_id, doc.file_path)
                #         weknora_doc_id = weknora_result.get("data", {}).get("id")
                        
                #         doc.weknora_knowledge_id = weknora_doc_id
                #         doc.processing_progress = 80
                #         logger.info(f"文档同步到 WeKnora 成功: {weknora_doc_id}")
                #     except Exception as e:
                #         logger.error(f"同步文档到 WeKnora 失败: {str(e)}")
                #         if hasattr(e, "response") and hasattr(e.response, "text"):
                #             logger.error(f"WeKnora 错误详情: {e.response.text}")
                        
                #         # 同步失败暂不标记为整个任务失败，但记录错误
                #         doc.error_message = f"WeKnora同步失败: {str(e)}"
                
                # 完成
                doc.status = "completed"
                doc.processing_progress = 100
                await session.commit()
                
                # 3. 索引到 Elasticsearch
                # 3.1 索引文档级信息 (用于传统搜索)
                await search_service.index_document(doc.id, {
                    "title": doc.title,
                    "content": doc.content,
                    "user_id": doc.user_id,
                    "knowledge_base_id": doc.knowledge_base_id,
                    "directory_id": doc.directory_id,
                    "document_type": doc.document_type,
                    "created_at": doc.created_at.isoformat(),
                    "updated_at": doc.updated_at.isoformat()
                })
                
                # 3.2 切片并向量化 (用于 RAG)
                if doc.content and doc.content.strip():
                    try:
                        doc.processing_progress = 60
                        await session.commit()
                        
                        logger.info(f"开始切分文档 {doc_id} 并生成向量...")
                        chunks_text = chunking_service.split_text(doc.content)
                        logger.info(f"文档 {doc_id} 切分完成，共 {len(chunks_text)} 个片段")
                        
                        chunks_data = []
                        for i, text in enumerate(chunks_text):
                            # 生成向量
                            vector = await llm_service.get_embeddings(text)
                            if not vector:
                                logger.warning(f"文档 {doc_id} 片段 {i} 向量生成失败，跳过")
                                continue
                                
                            chunks_data.append({
                                "doc_id": doc.id,
                                "chunk_index": i,
                                "content": text,
                                "vector": vector,
                                "knowledge_base_id": doc.knowledge_base_id,
                                "metadata": {
                                    "title": doc.title,
                                    "source": doc.file_path,
                                    "created_at": doc.created_at.isoformat()
                                }
                            })
                            
                        # 批量索引切片
                        if chunks_data:
                            await search_service.index_chunks(chunks_data)
                            logger.info(f"文档 {doc_id} 切片索引完成")
                            
                    except Exception as e:
                        logger.error(f"文档 {doc_id} 切片/向量化失败: {e}")
                        # 不阻断主流程，但记录日志
                
                logger.info(f"文档 {doc_id} 处理完成")

            except Exception as e:
                logger.error(f"后台处理文档异常: {e}")
                # 尝试更新错误状态
                try:
                    doc.status = "failed"
                    doc.error_message = f"系统错误: {str(e)}"
                    await session.commit()
                except:
                    pass

    async def list_documents(
        self, 
        skip: int, 
        limit: int, 
        category: Optional[str], 
        user_id: int,
        knowledge_base_id: Optional[int] = None,
        directory_id: Optional[int] = None
    ) -> List[dict]:
        """获取文档列表"""
        query = select(Document).where(Document.user_id == user_id)
        
        if category:
            query = query.where(Document.document_type == category)
            
        if knowledge_base_id is not None:
            query = query.where(Document.knowledge_base_id == knowledge_base_id)
            
        if directory_id is not None:
            query = query.where(Document.directory_id == directory_id)
        # 如果指定了知识库但没有指定目录，且不是查询特定分类，通常只显示根目录下的文件
        # 或者根据业务需求显示所有。这里假设如果不传 directory_id，则显示该知识库下所有文件，或者只显示根目录
        # 现在的需求是 "目录结构也显示对应目录中有哪些内容"，这意味着点击目录只显示该目录内容
        # 如果 directory_id 为 None 但 knowledge_base_id 不为 None，我们可能想显示根目录(directory_id is NULL)
        # 但前端如果不传 directory_id，可能意味着 "所有" 或者 "根"。
        # 让我们约定：如果前端明确传了 directory_id (即使是 0 或特定值)，则过滤。
        # 如果没传 directory_id，但传了 knowledge_base_id，我们可能需要决定是显示所有还是仅根。
        # 通常文件管理器逻辑是：没选目录 -> 根目录。
        # 但为了灵活性，我们让前端控制。如果前端想看根目录，应该传 directory_id=0 或 null 且后端处理 logic。
        # 在 SQL 中，NULL 比较需要用 is_(None)。
        # 让我们修改逻辑：如果 knowledge_base_id 存在，且 directory_id 是 None，我们是否过滤 directory_id IS NULL?
        # 现在的 listKnowledge 实现中，如果不传 dirId，就是 undefined。
        # 为了兼容 "显示所有文档" 和 "显示根目录文档"，我们可以约定：
        # - directory_id=0 -> 根目录 (SQL directory_id IS NULL)
        # - directory_id=None -> 不过滤目录 (显示所有)
        # 但 Document 模型中 directory_id 是 ForeignKey，存的是 NULL。
        
        # 修正逻辑：严格按照传入参数过滤。
        # 前端需要配合：点击 "全部文档" -> directory_id=undefined (后端 None) -> 显示所有
        # 点击 "根目录" (如果有这个概念) -> directory_id=0 -> 后端转为 IS NULL
        
        # 实际上前端 KnowledgeTree 有 "全部文档" (selectedDirId=null)。
        # 当 selectedDirId=null 时，我们希望显示该知识库下的 *所有* 文档，还是 *根* 文档？
        # 通常 "全部文档" 意味着忽略目录层级。
        # 如果是点击了某个文件夹，则 selectedDirId=123。
        # 只有在特定 "根目录" 视图下才需要 IS NULL。
        # 目前前端逻辑是：selectedDirId === null -> "全部文档"。
        # 所以这里 directory_id=None 时不加过滤条件是正确的，显示该库所有文档。
            
        query = query.order_by(Document.created_at.desc()).offset(skip).limit(limit)
        result = await self.db.execute(query)
        docs = result.scalars().all()
        return [self._format_document(doc) for doc in docs]

    async def get_document(self, doc_id: int, user_id: int) -> dict:
        """获取文档详情 (本地数据库)"""
        query = select(Document).where(and_(Document.id == doc_id, Document.user_id == user_id))
        result = await self.db.execute(query)
        doc = result.scalars().first()
        if not doc:
            raise ValueError("文档不存在")

        summary = await self._get_latest_summary(doc.id, user_id)
        payload = self._format_document(doc)
        # 返回文档内容
        payload["content"] = doc.content
        
        if summary:
            payload["latest_summary"] = summary
        return payload

    async def update_document(self, doc_id: int, doc_data: Dict, user_id: int) -> dict:
        """更新文档信息 (含版本控制)"""
        query = select(Document).where(and_(Document.id == doc_id, Document.user_id == user_id)).options(selectinload(Document.tags))
        result = await self.db.execute(query)
        doc = result.scalars().first()
        if not doc:
            raise ValueError("文档不存在")

        if doc_data.get("title"):
            doc.title = doc_data["title"]
        if doc_data.get("category"):
            doc.document_type = doc_data["category"]
            
        # Handle content update
        new_content = doc_data.get("content")
        # Check if content is provided (not None) and different
        if new_content is not None and new_content != doc.content:
            doc.content = new_content
            # is_content_changed flag is used below
            is_content_changed = True
        else:
            is_content_changed = False

        if doc_data.get("review_status"):
            doc.review_status = doc_data["review_status"]
            
        doc.updated_at = datetime.utcnow()
        
        if is_content_changed:
            doc.current_version += 1
            # 创建新版本
            version = DocumentVersion(
                document_id=doc.id,
                version_number=doc.current_version,
                title=doc.title,
                content=doc.content,
                created_by=user_id,
                change_log=doc_data.get("change_log", f"更新版本 {doc.current_version}")
            )
            self.db.add(version)

        await self.db.commit()
        await self.db.refresh(doc)

        # 同步更新 ES
        await search_service.index_document(doc.id, {
            "title": doc.title,
            "content": doc.content,
            "user_id": doc.user_id,
            "knowledge_base_id": doc.knowledge_base_id,
            "directory_id": doc.directory_id,
            "document_type": doc.document_type,
            "tags": [t.name for t in doc.tags],
            "updated_at": doc.updated_at.isoformat()
        })

        # 同步更新 WeKnora 中的标题
        if doc.weknora_knowledge_id:
            try:
                await self.weknora.update_document(doc.weknora_knowledge_id, title=doc.title)
            except Exception as e:
                logger.error(f"同步更新 WeKnora 标题失败: {e}")

        return self._format_document(doc)

    async def delete_document(self, doc_id: int, user_id: int) -> None:
        """删除文档"""
        query = select(Document).where(and_(Document.id == doc_id, Document.user_id == user_id))
        result = await self.db.execute(query)
        doc = result.scalars().first()
        if not doc:
            raise ValueError("文档不存在")

        # 同步从 WeKnora 删除
        if doc.weknora_knowledge_id:
            try:
                await self.weknora.delete_document(doc.weknora_knowledge_id)
            except Exception as e:
                logger.error(f"同步删除 WeKnora 文档失败: {e}")

        # 同步从 ES 删除
        await search_service.delete_document(doc.id)

        await self.db.execute(delete(DocumentSummary).where(DocumentSummary.document_id == doc.id))
        await self.db.delete(doc)
        await self.db.commit()

    # =====================
    # WeKnora 整合逻辑
    # =====================

    async def search_similar(self, query: str, knowledge_base_ids: List[str] = None, limit: int = 10) -> List[dict]:
        """使用 Elasticsearch 进行全文搜索"""
        logger.info(f"使用 ES 搜索文献: {query}, KBs: {knowledge_base_ids}")
        
        filters = {}
        # 注意：ES 中 knowledge_base_id 是 integer，但请求参数可能是 string list
        if knowledge_base_ids:
            # 尝试转换 ID 类型
            kb_ids_int = []
            for kid in knowledge_base_ids:
                try:
                    kb_ids_int.append(int(kid))
                except:
                    pass
            if kb_ids_int:
                filters["knowledge_base_id"] = kb_ids_int
                
        # 也可以添加 directory_id 过滤，如果需要的话，目前 API 没传

        # 调用 search_service
        results = await search_service.search(query, filters, page=1, size=limit)
        
        # 格式化为前端期望的格式 (类似于 KnowledgeSearchChunk 或 SearchResult)
        # 前端 KnowledgeSearchChunk: { id, content, knowledge_title, score, ... }
        # ES 返回: { total, items: [{id, title, content, title_highlight, content_highlight, ...}] }
        
        formatted_items = []
        for item in results["items"]:
            # 优先使用高亮内容
            display_content = " ... ".join(item.get("content_highlight", [])) or item.get("content", "")[:200]
            display_title = item.get("title_highlight", item.get("title"))
            
            formatted_items.append({
                "id": str(item["id"]),
                "knowledge_id": str(item["id"]), # 兼容旧字段
                "knowledge_title": display_title,
                "content": display_content,
                "score": item["score"],
                "chunk_type": "document",
                "knowledge_source": item.get("document_type", "file"),
                "knowledge_filename": item.get("title")
            })
            
        return formatted_items

    async def get_document_details(self, doc_id: int) -> dict:
        """获取文档详情 (自动映射 ID)"""
        # 先查本地找到 WeKnora ID
        query = select(Document).where(Document.id == doc_id)
        result = await self.db.execute(query)
        doc = result.scalars().first()
        
        if doc and doc.weknora_knowledge_id:
            return await self.weknora.get_document(doc.weknora_knowledge_id)
        
        raise ValueError("文档未同步到 WeKnora 或不存在")

    async def summarize_document(self, doc_id: int, user_id: int) -> dict:
        """使用 WeKnora 内容生成文档摘要并持久化"""
        logger.info(f"开始为文档生成摘要: {doc_id}")
        
        query = select(Document).where(Document.id == doc_id)
        result = await self.db.execute(query)
        doc = result.scalars().first()
        
        if not doc:
            return {"success": False, "message": "文档不存在"}

        content = ""
        if doc.weknora_knowledge_id:
            content = await self.weknora.get_document_full_content(doc.weknora_knowledge_id)
        
        if not content:
            content = doc.content
        
        if not content:
            return {"success": False, "message": "未找到文档内容"}
        
        summary_text = await llm_service.generate_document_summary(content)
        
        # --- 持久化摘要到数据库 ---
        summary = DocumentSummary(
            document_id=doc.id,
            user_id=user_id,
            summary_level="paragraph",
            summary_text=summary_text,
            model_name=llm_service.model
        )
        self.db.add(summary)
        await self.db.commit()
        await self.db.refresh(summary)

        return {"success": True, "data": self._format_summary(summary)}

    async def get_document_concepts(self, doc_id: int) -> dict:
        """从 WeKnora 内容获取文档关键概念"""
        logger.info(f"开始提取文档关键概念: {doc_id}")
        
        query = select(Document).where(Document.id == doc_id)
        result = await self.db.execute(query)
        doc = result.scalars().first()
        
        if not doc: return {"success": False, "message": "文档不存在"}

        content = ""
        if doc.weknora_knowledge_id:
            content = await self.weknora.get_document_full_content(doc.weknora_knowledge_id)
        
        if not content: content = doc.content
        
        if not content: return {"success": False, "message": "未找到文档内容"}
        
        concepts = await llm_service.extract_document_concepts(content)
        return {"success": True, "data": concepts}

    async def get_document_citations(self, doc_id: int) -> dict:
        """从 WeKnora 内容获取文档引用关系"""
        logger.info(f"开始提取文档引用关系: {doc_id}")
        
        query = select(Document).where(Document.id == doc_id)
        result = await self.db.execute(query)
        doc = result.scalars().first()
        
        if not doc: return {"success": False, "message": "文档不存在"}

        content = ""
        if doc.weknora_knowledge_id:
            content = await self.weknora.get_document_full_content(doc.weknora_knowledge_id)
        
        if not content: content = doc.content
        
        if not content: return {"success": False, "message": "未找到文档内容"}
        
        citations = await llm_service.extract_document_citations(content)
        return {"success": True, "data": citations}

    # =====================
    # 辅助方法
    # =====================

    def _parse_file(self, path: str, ext: str) -> str:
        if ext == "pdf":
            return self._parse_pdf(path)
        if ext in {"txt", "md"}:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        if ext == "docx":
            doc = DocxDocument(path)
            return "\n".join([p.text for p in doc.paragraphs if p.text])
        return ""

    def _parse_pdf(self, path: str) -> str:
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                texts.append(page.extract_text() or "")
        return "\n".join(texts)

    def _strip_html(self, html: str) -> str:
        html = re.sub(r"<script.*?>.*?</script>", "", html, flags=re.S | re.I)
        html = re.sub(r"<style.*?>.*?</style>", "", html, flags=re.S | re.I)
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    async def _get_latest_summary(self, doc_id: int, user_id: int) -> Optional[dict]:
        query = (
            select(DocumentSummary)
            .where(and_(DocumentSummary.document_id == doc_id, DocumentSummary.user_id == user_id))
            .order_by(DocumentSummary.created_at.desc())
            .limit(1)
        )
        result = await self.db.execute(query)
        summary = result.scalars().first()
        if not summary:
            return None
        return self._format_summary(summary)

    def _format_document(self, doc: Document) -> dict:
        return {
            "id": doc.id,
            "title": doc.title,
            "document_type": doc.document_type,
            "source_type": doc.source_type,
            "source_url": doc.source_url,
            "file_path": doc.file_path,
            "weknora_knowledge_id": doc.weknora_knowledge_id,
            "weknora_kb_id": doc.weknora_kb_id,
            "status": doc.status,
            "processing_progress": doc.processing_progress,
            "error_message": doc.error_message,
            "created_at": doc.created_at,
            "updated_at": doc.updated_at,
        }

    def _format_summary(self, summary: DocumentSummary) -> dict:
        return {
            "id": summary.id,
            "document_id": summary.document_id,
            "summary_level": summary.summary_level,
            "summary_text": summary.summary_text,
            "quality_score": summary.quality_score,
            "model_name": summary.model_name,
            "created_at": summary.created_at,
        }
