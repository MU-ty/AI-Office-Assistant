"""
文档生成服务 - 支持多种格式输出
生成Markdown、PDF、Word等格式的会议纪要

表格中标黄部分 - PDF生成、Word文档生成、Markdown生成
"""

from typing import Dict, List, Optional, Any
from datetime import datetime
import json
from html import escape
from app.utils.logger import get_logger

logger = get_logger(__name__)


class DocumentGenerationService:
    """文档生成服务 - 支持多种格式"""
    
    # ============================================================
    # Markdown生成
    # ============================================================
    
    def generate_markdown(
        self,
        title: str,
        meeting_data: Dict,
        include_toc: bool = True
    ) -> str:
        """
        生成Markdown格式的会议纪要
        
        Args:
            title: 纪要标题
            meeting_data: 会议数据，包含：
                - date: 会议日期
                - participants: 参与人列表
                - agendas: 议程列表
                - key_points: 关键点
                - decisions: 决议列表
                - action_items: Action Items列表
                - transcription: 转录文本
            include_toc: 是否生成目录
            
        Returns:
            Markdown格式文本
        """
        md_lines = []
        
        # 标题
        md_lines.append(f"# {title}")
        md_lines.append("")
        
        # 基本信息
        md_lines.append("## 基本信息")
        md_lines.append("")
        if 'date' in meeting_data:
            md_lines.append(f"- **会议日期**: {meeting_data['date']}")
        if 'participants' in meeting_data:
            participants = ", ".join(meeting_data['participants']) if isinstance(meeting_data['participants'], list) else str(meeting_data['participants'])
            md_lines.append(f"- **参与人**: {participants}")
        md_lines.append("")
        
        # 目录（可选）
        if include_toc:
            md_lines.append("## 目录")
            md_lines.append("1. [议程](#议程)")
            md_lines.append("2. [关键点](#关键点)")
            md_lines.append("3. [决议](#决议)")
            md_lines.append("4. [Action Items](#action-items)")
            md_lines.append("")
        
        # 议程
        if 'agendas' in meeting_data and meeting_data['agendas']:
            md_lines.append("## 议程")
            md_lines.append("")
            for i, agenda in enumerate(meeting_data['agendas'], 1):
                if isinstance(agenda, dict):
                    md_lines.append(f"### {i}. {agenda.get('title', '议题' + str(i))}")
                    if 'description' in agenda:
                        md_lines.append(f"{agenda['description']}")
                else:
                    md_lines.append(f"### {i}. {agenda}")
                md_lines.append("")
        
        # 关键点
        if 'key_points' in meeting_data and meeting_data['key_points']:
            md_lines.append("## 关键点")
            md_lines.append("")
            for point in meeting_data['key_points']:
                md_lines.append(f"- {point}")
            md_lines.append("")
        
        # 决议
        if 'decisions' in meeting_data and meeting_data['decisions']:
            md_lines.append("## 决议")
            md_lines.append("")
            for i, decision in enumerate(meeting_data['decisions'], 1):
                if isinstance(decision, dict):
                    md_lines.append(f"### 决议{i}")
                    md_lines.append(f"{decision.get('content', '决议内容')}")
                else:
                    md_lines.append(f"### 决议{i}")
                    md_lines.append(f"{decision}")
                md_lines.append("")
        
        # Action Items
        if 'action_items' in meeting_data and meeting_data['action_items']:
            md_lines.append("## Action Items")
            md_lines.append("")
            for i, item in enumerate(meeting_data['action_items'], 1):
                if isinstance(item, dict):
                    owner = item.get('owner', '待定')
                    due_date = item.get('due_date', '待定')
                    content = item.get('content', 'Action Item' + str(i))
                    md_lines.append(f"- [ ] **{content}** (负责人: {owner}, 截止日期: {due_date})")
                else:
                    md_lines.append(f"- [ ] {item}")
            md_lines.append("")
        
        # 转录文本（可选）
        if 'transcription' in meeting_data and meeting_data['transcription']:
            md_lines.append("## 完整转录")
            md_lines.append("")
            md_lines.append("```")
            md_lines.append(meeting_data['transcription'])
            md_lines.append("```")
            md_lines.append("")
        
        # 页脚
        md_lines.append("---")
        md_lines.append(f"*生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
        
        return "\n".join(md_lines)
    
    # ============================================================
    # PDF生成（使用reportlab）
    # ============================================================
    
    def generate_pdf(
        self,
        title: str,
        meeting_data: Dict,
        output_path: str
    ) -> bool:
        """
        生成PDF格式的会议纪要
        
        Args:
            title: 纪要标题
            meeting_data: 会议数据
            output_path: 输出文件路径
            
        Returns:
            成功返回True，失败返回False
        """
        try:
            from reportlab.lib.pagesizes import A4, inch  # type: ignore
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
            from reportlab.lib.units import cm  # type: ignore
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak  # type: ignore
            from reportlab.lib import colors  # type: ignore
            from reportlab.pdfbase import pdfmetrics  # type: ignore
            from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
            import os
            import platform
            
            # 注册中文字体
            font_registered = False
            system = platform.system()
            font_name = 'SimHei'
            
            # 根据不同系统尝试不同的字体路径
            font_paths = []
            if system == 'Windows':
                font_paths = [
                    'C:\\Windows\\Fonts\\simhei.ttf',
                    'C:\\Windows\\Fonts\\SimHei.ttf',
                    'C:\\Windows\\Fonts\\msyh.ttf',  # Microsoft YaHei
                    'C:\\Windows\\Fonts\\msyhbd.ttf', # Microsoft YaHei Bold
                    'C:\\winnt\\Fonts\\simhei.ttf',
                ]
            elif system == 'Darwin':  # macOS
                font_paths = [
                    '/Library/Fonts/SimHei.ttf',
                    '/System/Library/Fonts/STHeiti Medium.ttc',
                    '/System/Library/Fonts/PingFang.ttc',
                ]
            else:  # Linux
                font_paths = [
                    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
                    '/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf',
                    '/usr/share/fonts/wqy-microhei/wqy-microhei.ttc',
                ]
            
            # 尝试加载字体
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        logger.info(f"成功加载字体: {font_path}")
                        font_registered = True
                        break
                    except Exception as e:
                        logger.debug(f"加载字体失败 {font_path}: {e}")
                        continue
            
            if not font_registered:
                logger.warning("未能找到中文字体，PDF中文显示可能有问题。建议安装SimHei字体")
                # 使用Helvetica作为fallback
                font_name = 'Helvetica'
            
            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            elements = []
            
            # 样式定义 - 明确指定中文字体
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=font_name,
                fontSize=18,
                textColor=colors.HexColor('#1f4788'),
                spaceAfter=30,
                alignment=1  # 居中
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=font_name,
                fontSize=14,
                textColor=colors.HexColor('#2e5c8a'),
                spaceAfter=12,
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
            )

            def safe_paragraph(text, style):
                """Helper to safely create paragraphs with escaped text"""
                if not text:
                    return Paragraph("", style)
                # Escape HTML/XML chars
                safe_text = escape(str(text))
                # Convert newlines to <br/>
                safe_text = safe_text.replace('\n', '<br/>')
                return Paragraph(safe_text, style)
            
            # 添加标题
            elements.append(safe_paragraph(title, title_style))
            elements.append(Spacer(1, 0.3*inch))
            
            # 添加基本信息表
            if 'date' in meeting_data or 'participants' in meeting_data:
                info_data = [['字段', '内容']]
                if 'date' in meeting_data:
                    info_data.append(['会议日期', str(meeting_data['date'])])
                if 'participants' in meeting_data:
                    participants = ", ".join(meeting_data['participants']) if isinstance(meeting_data['participants'], list) else str(meeting_data['participants'])
                    info_data.append(['参与人', participants])
                
                info_table = Table(info_data, colWidths=[2*cm, 14*cm])
                info_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2e5c8a')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), font_name),
                    ('FONTNAME', (0, 1), (-1, -1), font_name),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('FONTSIZE', (0, 1), (-1, -1), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ]))
                elements.append(info_table)
                elements.append(Spacer(1, 0.3*inch))
            
            # 添加各个部分
            if 'agendas' in meeting_data and meeting_data['agendas']:
                elements.append(Paragraph('议程', heading_style))
                for i, agenda in enumerate(meeting_data['agendas'], 1):
                    agenda_text = f"{i}. {agenda if isinstance(agenda, str) else agenda.get('title', '')}"
                    elements.append(safe_paragraph(agenda_text, normal_style))
                elements.append(Spacer(1, 0.2*inch))
            
            # 添加段落内容（从minutes文本）
            if 'paragraphs' in meeting_data and meeting_data['paragraphs']:
                elements.append(Paragraph('会议内容', heading_style))
                for para in meeting_data['paragraphs']:
                    if para:
                        # 如果段落太长，自动换行
                        elements.append(safe_paragraph(para, normal_style))
                        elements.append(Spacer(1, 0.1*inch))
                elements.append(Spacer(1, 0.2*inch))
            
            # 添加关键点
            if 'key_points' in meeting_data and meeting_data['key_points']:
                elements.append(Paragraph('关键点', heading_style))
                for point in meeting_data['key_points']:
                    elements.append(safe_paragraph(f"• {point}", normal_style))
                elements.append(Spacer(1, 0.2*inch))
            
            # 添加决议
            if 'decisions' in meeting_data and meeting_data['decisions']:
                elements.append(Paragraph('决议', heading_style))
                for i, decision in enumerate(meeting_data['decisions'], 1):
                    decision_text = f"{i}. {decision if isinstance(decision, str) else decision.get('content', '')}"
                    elements.append(safe_paragraph(decision_text, normal_style))
                elements.append(Spacer(1, 0.2*inch))
            
            # 添加Action Items
            if 'action_items' in meeting_data and meeting_data['action_items']:
                elements.append(Paragraph('Action Items', heading_style))
                for item in meeting_data['action_items']:
                    if isinstance(item, dict):
                        item_text = f"• {item.get('content', '')} (负责人: {item.get('owner', '')}, 截止日期: {item.get('due_date', '')})"
                    else:
                        item_text = f"• {item}"
                    elements.append(safe_paragraph(item_text, normal_style))
                elements.append(Spacer(1, 0.2*inch))
            
            # 生成PDF
            doc.build(elements)
            logger.info(f"PDF生成成功: {output_path}")
            return True
            
        except ImportError as e:
            logger.error(f"reportlab 未安装: {e}，请运行: pip install reportlab")
            return False
        except Exception as e:
            logger.error(f"PDF生成失败: {type(e).__name__}: {e}", exc_info=True)
            return False
    
    # ============================================================
    # Word文档生成（使用python-docx）
    # ============================================================
    
    def generate_docx(
        self,
        title: str,
        meeting_data: Dict,
        output_path: str
    ) -> bool:
        """
        生成Word格式的会议纪要
        
        Args:
            title: 纪要标题
            meeting_data: 会议数据
            output_path: 输出文件路径
            
        Returns:
            成功返回True，失败返回False
        """
        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt, RGBColor, Inches  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
            
            # 创建Document对象
            doc = Document()
            
            # 设置中文字体
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)
            
            # 添加标题
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.color.rgb = RGBColor(31, 71, 136)
            
            # 添加基本信息
            doc.add_heading('基本信息', level=2)
            if 'date' in meeting_data:
                doc.add_paragraph(f"会议日期：{meeting_data['date']}")
            if 'participants' in meeting_data:
                participants = ", ".join(meeting_data['participants']) if isinstance(meeting_data['participants'], list) else str(meeting_data['participants'])
                doc.add_paragraph(f"参与人：{participants}")
            
            # 添加议程
            if 'agendas' in meeting_data and meeting_data['agendas']:
                doc.add_heading('议程', level=2)
                for i, agenda in enumerate(meeting_data['agendas'], 1):
                    agenda_text = f"{i}. {agenda if isinstance(agenda, str) else agenda.get('title', '')}"
                    doc.add_paragraph(agenda_text, style='List Number')
            
            # 添加段落内容（从minutes文本）
            if 'paragraphs' in meeting_data and meeting_data['paragraphs']:
                doc.add_heading('会议内容', level=2)
                for para in meeting_data['paragraphs']:
                    if para:
                        doc.add_paragraph(str(para))
            
            # 添加关键点
            if 'key_points' in meeting_data and meeting_data['key_points']:
                doc.add_heading('关键点', level=2)
                for point in meeting_data['key_points']:
                    doc.add_paragraph(point, style='List Bullet')
            
            # 添加决议
            if 'decisions' in meeting_data and meeting_data['decisions']:
                doc.add_heading('决议', level=2)
                for i, decision in enumerate(meeting_data['decisions'], 1):
                    decision_text = f"{i}. {decision if isinstance(decision, str) else decision.get('content', '')}"
                    doc.add_paragraph(decision_text, style='List Number')
            
            # 添加Action Items
            if 'action_items' in meeting_data and meeting_data['action_items']:
                doc.add_heading('Action Items', level=2)
                for item in meeting_data['action_items']:
                    if isinstance(item, dict):
                        item_text = f"{item.get('content', '')} (负责人: {item.get('owner', '')}, 截止日期: {item.get('due_date', '')})"
                    else:
                        item_text = str(item)
                    doc.add_paragraph(item_text, style='List Bullet')
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"Word文档生成成功: {output_path}")
            return True
            
        except ImportError:
            logger.error("python-docx 未安装，请运行: pip install python-docx")
            return False
        except Exception as e:
            logger.error(f"Word文档生成失败: {type(e).__name__}: {e}", exc_info=True)
            return False
    
    # ============================================================
    # JSON生成
    # ============================================================
    
    def generate_json(self, meeting_data: Dict) -> str:
        """
        生成JSON格式的会议纪要
        
        Args:
            meeting_data: 会议数据
            
        Returns:
            JSON格式字符串
        """
        return json.dumps(meeting_data, ensure_ascii=False, indent=2)


# 全局实例
document_generation_service = DocumentGenerationService()
